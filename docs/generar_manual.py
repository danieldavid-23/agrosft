"""
Generador del Manual de Instalacion - AgroSFT
Ejecutar: python docs/generar_manual.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ============================================================
# ESTILOS GLOBALES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # Verde oscuro
    hs.font.name = 'Calibri'

# ============================================================
# PORTADILLA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MANUAL DE INSTALACION')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AgroSFT')
r.bold = True
r.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema de Gestion Agricola')
r.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 1.0')
r.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SENA - Regional Antioquia - Medellin')
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
hoy = datetime.date.today()
r = p.add_run(f'{hoy.strftime("%B").capitalize()} de {hoy.year}')
r.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# LISTA DE COLABORADORES
# ============================================================
doc.add_heading('Lista de Colaboradores', level=1)

tabla = doc.add_table(rows=5, cols=2)
tabla.style = 'Light Grid Accent 1'
tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

datos = [
    ('Rol', 'Nombre'),
    ('Equipo de Desarrollo', '[Nombres completos de los aprendices]'),
    ('Instructor(a)', '[Nombre completo del instructor]'),
    ('Responsable de Instalacion', '[Nombre y cargo]'),
    ('Cliente / Entidad', 'SENA - Servicio Nacional de Aprendizaje'),
]
for i, (c1, c2) in enumerate(datos):
    tabla.rows[i].cells[0].text = c1
    tabla.rows[i].cells[1].text = c2
    if i == 0:
        for cell in tabla.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDO
# ============================================================
doc.add_heading('Tabla de Contenido', level=1)
doc.add_paragraph('[Actualizar esta tabla con las paginas correspondientes al finalizar el documento]')
doc.add_page_break()

# ============================================================
# TABLA DE IMAGENES
# ============================================================
doc.add_heading('Tabla de Imagenes', level=1)
doc.add_paragraph('[Incluir esta tabla solo si se usan capturas, diagramas o evidencias numeradas]')
doc.add_page_break()

# ============================================================
# 1. INTRODUCCION
# ============================================================
doc.add_heading('1. Introduccion', level=1)

doc.add_heading('1.1 Objetivo', level=2)
doc.add_paragraph(
    'El presente manual describe el procedimiento completo para la instalacion, '
    'configuracion y ejecucion del sistema AgroSFT en un ambiente de desarrollo local. '
    'Esta dirigido a aprendices, instructores y tecnicos responsables de desplegar '
    'el sistema en estaciones de trabajo individuales.'
)

doc.add_heading('1.2 Alcance', level=2)
doc.add_paragraph(
    'Este manual cubre la instalacion desde el repositorio fuente hasta la verificacion '
    'basica del sistema funcionando. No incluye despliegue en servidor de produccion, '
    'configuracion de dominio ni certificados SSL.'
)

doc.add_heading('1.3 Audiencia', level=2)
doc.add_paragraph(
    'Aprendices del programa de formacion en Desarrollo de Software del SENA, '
    'instructores y personal tecnico con conocimientos basicos de Python, terminal '
    'de comandos y bases de datos.'
)

doc.add_heading('1.4 Version del Software', level=2)
tabla = doc.add_table(rows=4, cols=2)
tabla.style = 'Light Grid Accent 1'
for i, (c1, c2) in enumerate([
    ('Componente', 'Version'),
    ('AgroSFT', '1.0'),
    ('Django', '5.0.14'),
    ('Vue.js', '3.5+'),
]):
    tabla.rows[i].cells[0].text = c1
    tabla.rows[i].cells[1].text = c2
    if i == 0:
        for cell in tabla.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ============================================================
# 2. REQUERIMIENTOS
# ============================================================
doc.add_heading('2. Requerimientos', level=1)

doc.add_heading('2.1 Requerimientos de Hardware', level=2)
doc.add_paragraph('Minimos:')
doc.add_paragraph('Procesador: 2 nucleos (Intel i3 o equivalente)', style='List Bullet')
doc.add_paragraph('Memoria RAM: 4 GB', style='List Bullet')
doc.add_paragraph('Almacenamiento: 2 GB de espacio libre en disco', style='List Bullet')
doc.add_paragraph('Resolucion: 1024x768', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Recomendados:')
doc.add_paragraph('Procesador: 4 nucleos (Intel i5 o equivalente)', style='List Bullet')
doc.add_paragraph('Memoria RAM: 8 GB', style='List Bullet')
doc.add_paragraph('Almacenamiento: 5 GB de espacio libre en disco', style='List Bullet')
doc.add_paragraph('Resolucion: 1920x1080', style='List Bullet')

doc.add_heading('2.2 Requerimientos de Software', level=2)
tabla = doc.add_table(rows=8, cols=3)
tabla.style = 'Light Grid Accent 1'
software = [
    ('Componente', 'Version Minima', 'Version Recomendada'),
    ('Sistema Operativo', 'Windows 10 / Ubuntu 20.04', 'Windows 11 / Ubuntu 22.04'),
    ('Python', '3.10', '3.12+'),
    ('Node.js', '18.0', '20 LTS'),
    ('MySQL / MariaDB', '10.6', '10.11 / 11.x'),
    ('Git', '2.30', '2.40+'),
    ('pip', '22.0', 'Ultima estable'),
    ('Navegador Web', 'Chrome 90 / Firefox 90', 'Ultima version estable'),
]
for i, row_data in enumerate(software):
    for j, val in enumerate(row_data):
        tabla.rows[i].cells[j].text = val
        if i == 0:
            for p in tabla.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('2.3 Requerimientos de Red y Permisos', level=2)
doc.add_paragraph('Puertos requeridos:')
doc.add_paragraph('Puerto 8000: Servidor de desarrollo Django (acceso local)', style='List Bullet')
doc.add_paragraph('Puerto 3306: Servidor MySQL/MariaDB (acceso local)', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Permisos:')
doc.add_paragraph('Acceso de lectura/escritura al directorio del proyecto', style='List Bullet')
doc.add_paragraph('Permisos de administrador en MySQL para crear la base de datos', style='List Bullet')
doc.add_paragraph('Conexion a internet para instalar dependencias y autenticacion Google OAuth', style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. GUIA DE INSTALACION
# ============================================================
doc.add_heading('3. Guia de Instalacion', level=1)

doc.add_heading('3.1 Preparacion del Ambiente', level=2)
doc.add_paragraph(
    'Antes de iniciar la instalacion, verifique que tenga instalados Python, '
    'Node.js, MySQL/MariaDB y Git. Abra una terminal y ejecute los siguientes '
    'comandos para confirmar:'
)
doc.add_paragraph('python --version', style='List Bullet')
doc.add_paragraph('node --version', style='List Bullet')
doc.add_paragraph('mysql --version', style='List Bullet')
doc.add_paragraph('git --version', style='List Bullet')

doc.add_heading('3.2 Instalacion', level=2)

doc.add_heading('Paso 1: Clonar el repositorio', level=3)
doc.add_paragraph('git clone [URL_DEL_REPOSITORIO]')
doc.add_paragraph('cd agrosft')

doc.add_heading('Paso 2: Crear entorno virtual de Python', level=3)
doc.add_paragraph('python -m venv venv')
doc.add_paragraph()
doc.add_paragraph('Activar el entorno virtual:')
doc.add_paragraph('Windows: venv\\Scripts\\activate', style='List Bullet')
doc.add_paragraph('Linux/Mac: source venv/bin/activate', style='List Bullet')

doc.add_heading('Paso 3: Instalar dependencias de Python', level=3)
doc.add_paragraph('pip install -r requirements.txt')

doc.add_heading('Paso 4: Instalar dependencias de Node.js (frontend)', level=3)
doc.add_paragraph('npm install')
doc.add_paragraph('npm run build')

doc.add_heading('Paso 5: Configurar variables de entorno', level=3)
doc.add_paragraph(
    'Cree o edite el archivo .env en la raiz del proyecto con las siguientes variables:'
)
doc.add_paragraph('SECRET_KEY=[Clave secreta de Django]', style='List Bullet')
doc.add_paragraph('DEBUG=True', style='List Bullet')
doc.add_paragraph('DB_NAME=agrosft', style='List Bullet')
doc.add_paragraph('DB_USER=root', style='List Bullet')
doc.add_paragraph('DB_PASSWORD=[Su contrasena de MySQL]', style='List Bullet')
doc.add_paragraph('DB_HOST=127.0.0.1', style='List Bullet')
doc.add_paragraph('DB_PORT=3306', style='List Bullet')
doc.add_paragraph('GOOGLE_CLIENT_ID=[ID de Google OAuth]', style='List Bullet')
doc.add_paragraph('GOOGLE_CLIENT_SECRET=[Secreto de Google OAuth]', style='List Bullet')
doc.add_paragraph('BREVO_API_KEY=[API Key de Brevo para correos]', style='List Bullet')
doc.add_paragraph('DEFAULT_FROM_EMAIL=agrosft.soporte@gmail.com', style='List Bullet')

doc.add_heading('Paso 6: Crear la base de datos', level=3)
doc.add_paragraph('Acceda a MySQL y cree la base de datos:')
doc.add_paragraph('mysql -u root -p')
doc.add_paragraph("CREATE DATABASE agrosft CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;")

doc.add_heading('Paso 7: Ejecutar migraciones y configurar esquema', level=3)
doc.add_paragraph(
    'El proyecto utiliza modelos con managed=False, por lo que el esquema se gestiona '
    'externamente. Ejecute las migraciones de Django para las tablas internas:'
)
doc.add_paragraph('python manage.py migrate')
doc.add_paragraph()
doc.add_paragraph(
    'Importe el esquema de las tablas principales desde phpMyAdmin o usando los scripts SQL '
    'del directorio scripts/. Los triggers de stock se instalan con:'
)
doc.add_paragraph('mysql -u root -p agrosft < scripts/trigger_proteccion_stock.sql')
doc.add_paragraph('mysql -u root -p agrosft < scripts/trigger_stock_vendida.sql')

doc.add_heading('Paso 8: Crear superusuario', level=3)
doc.add_paragraph('python crear_superusuario.py')
doc.add_paragraph()
doc.add_paragraph('O manualmente:')
doc.add_paragraph('python manage.py createsuperuser')

doc.add_heading('Paso 9: Verificar tipos de movimiento', level=3)
doc.add_paragraph('python scripts/asegurar_tipos_movimiento.py')

doc.add_heading('Paso 10: Ejecutar el servidor', level=3)
doc.add_paragraph('python manage.py runserver')
doc.add_paragraph()
doc.add_paragraph('Acceda al sistema en: http://127.0.0.1:8000')

doc.add_page_break()

# ============================================================
# 4. DESINSTALACION
# ============================================================
doc.add_heading('4. Desinstalacion', level=1)
doc.add_paragraph('Para desinstalar el sistema:')
doc.add_paragraph('1. Detenga el servidor de desarrollo (Ctrl+C en la terminal).', style='List Bullet')
doc.add_paragraph('2. Desactive el entorno virtual: deactivate', style='List Bullet')
doc.add_paragraph('3. Elimine el directorio del proyecto:', style='List Bullet')
doc.add_paragraph('   rmdir /s /q agrosft  (Windows)')
doc.add_paragraph('   rm -rf agrosft  (Linux/Mac)')
doc.add_paragraph()
doc.add_paragraph(
    'Para conservar los datos, NO elimine la base de datos agrosft en MySQL. '
    'Para eliminarla completamente:'
)
doc.add_paragraph('mysql -u root -p -e "DROP DATABASE agrosft;"')

doc.add_page_break()

# ============================================================
# 5. CONFIGURACION
# ============================================================
doc.add_heading('5. Configuracion', level=1)

doc.add_heading('5.1 Archivo .env', level=2)
doc.add_paragraph('Variables principales del archivo .env:')
tabla = doc.add_table(rows=9, cols=3)
tabla.style = 'Light Grid Accent 1'
env_data = [
    ('Variable', 'Descripcion', 'Valor por defecto'),
    ('SECRET_KEY', 'Clave secreta de Django', 'django-insecure-dev-only...'),
    ('DEBUG', 'Modo depuracion (True/False)', 'True'),
    ('DB_NAME', 'Nombre de la base de datos', 'agrosft'),
    ('DB_USER', 'Usuario de MySQL', 'root'),
    ('DB_PASSWORD', 'Contrasena de MySQL', '(vacio)'),
    ('GOOGLE_CLIENT_ID', 'ID de cliente Google OAuth', '(requerido para login Google)'),
    ('BREVO_API_KEY', 'API Key de Brevo para correos', '(requerido para reset de contrasena)'),
    ('DEFAULT_FROM_EMAIL', 'Correo de envio', 'agrosft.soporte@gmail.com'),
]
for i, row_data in enumerate(env_data):
    for j, val in enumerate(row_data):
        tabla.rows[i].cells[j].text = val
        if i == 0:
            for p in tabla.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('5.2 Configuracion de la Base de Datos', level=2)
doc.add_paragraph(
    'El archivo config/settings.py contiene la configuracion de conexion a MySQL. '
    'Los modelos de las apps usuarios, inventario, ventas y clientes utilizan '
    'managed=False, lo que significa que Django NO gestiona sus migraciones. '
    'El esquema se administra directamente en MySQL.'
)

doc.add_heading('5.3 Autenticacion', level=2)
doc.add_paragraph('El sistema soporta dos metodos de autenticacion:')
doc.add_paragraph('Inicio de sesion tradicional: correo + contrasena via el backend personalizado TblusuariosAuthBackend.', style='List Bullet')
doc.add_paragraph('Google OAuth2: Inicio de sesion con cuenta de Google. Requiere configurar GOOGLE_CLIENT_ID y GOOGLE_CLIENT_SECRET en .env.', style='List Bullet')

doc.add_heading('5.4 Sesion y Seguridad', level=2)
doc.add_paragraph('Sesiones en cookies firmadas (no usa tabla de sesiones en BD).', style='List Bullet')
doc.add_paragraph('Expiracion de sesion: 30 minutos o al cerrar el navegador.', style='List Bullet')
doc.add_paragraph('Headers de seguridad: XSS filter, Content-Type nosniff, X-Frame-Options DENY.', style='List Bullet')

doc.add_page_break()

# ============================================================
# 6. EJECUCION Y VERIFICACION
# ============================================================
doc.add_heading('6. Ejecucion y Verificaccion', level=1)

doc.add_heading('6.1 Iniciar el Sistema', level=2)
doc.add_paragraph('1. Active el entorno virtual: venv\\Scripts\\activate', style='List Bullet')
doc.add_paragraph('2. Ejecute: python manage.py runserver', style='List Bullet')
doc.add_paragraph('3. Abra el navegador en: http://127.0.0.1:8000', style='List Bullet')

doc.add_heading('6.2 Verificacion Basica', level=2)
doc.add_paragraph('Al acceder al sistema, verifique:')
doc.add_paragraph('Se muestra la pagina de inicio de sesion o el marketplace segun el estado de autenticacion.', style='List Bullet')
doc.add_paragraph('El login con correo y contrasena funciona correctamente.', style='List Bullet')
desc = doc.add_paragraph('El marketplace muestra los productos disponibles.', style='List Bullet')
doc.add_paragraph('El registro de usuarios crea un nuevo perfil en la base de datos.', style='List Bullet')
doc.add_paragraph('El carrito de compras permite agregar productos.', style='List Bullet')
doc.add_paragraph('La generacion de facturas en PDF funciona desde el historial.', style='List Bullet')

doc.add_heading('6.3 Panel de Administracion', level=2)
doc.add_paragraph('Acceda al panel de administracion en: http://127.0.0.1:8000/admin/')
doc.add_paragraph('Desde aqui puede gestionar usuarios, productos, categorias y revisar el registro de auditoria.')

doc.add_page_break()

# ============================================================
# 7. LICENCIAMIENTO Y GARANTIAS
# ============================================================
doc.add_heading('7. Licenciamiento y Garantias', level=1)

doc.add_heading('7.1 Licencias del Software', level=2)
tabla = doc.add_table(rows=8, cols=3)
tabla.style = 'Light Grid Accent 1'
lic_data = [
    ('Componente', 'Licencia', 'Tipo'),
    ('Django 5.0', 'BSD', 'Open Source'),
    ('Vue.js 3', 'MIT', 'Open Source'),
    ('Bootstrap 5.1', 'MIT', 'Open Source'),
    ('Font Awesome 6.4', 'CC BY 4.0 / SIL OFL', 'Open Source'),
    ('Vite 6', 'MIT', 'Open Source'),
    ('ReportLab 4.5', 'BSD', 'Open Source'),
    ('xhtml2pdf 0.2', 'GNU LGPL', 'Open Source'),
]
for i, row_data in enumerate(lic_data):
    for j, val in enumerate(row_data):
        tabla.rows[i].cells[j].text = val
        if i == 0:
            for p in tabla.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('7.2 Garantia', level=2)
doc.add_paragraph(
    'El software se proporciona "tal cual", sin garantia de ningun tipo, expresa o '
    'implimita. El equipo de desarrollo no asume responsabilidad por danos directos, '
    'indirectos o consecuentes derivados del uso del software.'
)

doc.add_page_break()

# ============================================================
# 8. ASISTENCIA TECNICA
# ============================================================
doc.add_heading('8. Asistencia Tecnica', level=1)
doc.add_paragraph('Canales de contacto para soporte tecnico:')
tabla = doc.add_table(rows=4, cols=2)
tabla.style = 'Light Grid Accent 1'
soporte = [
    ('Canal', 'Detalle'),
    ('Correo electronico', '[correo de soporte]'),
    ('Telefono / WhatsApp', '[numero de contacto]'),
    ('Horario', 'Lunes a Viernes, 8:00 a.m. - 5:00 p.m.'),
]
for i, (c1, c2) in enumerate(soporte):
    tabla.rows[i].cells[0].text = c1
    tabla.rows[i].cells[1].text = c2
    if i == 0:
        for cell in tabla.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()
doc.add_paragraph('Niveles de escalamiento:')
doc.add_paragraph('Nivel 1: Soporte basico (instalacion, configuracion, uso del sistema).', style='List Bullet')
doc.add_paragraph('Nivel 2: Soporte tecnico (errores de base de datos, triggers, integraciones).', style='List Bullet')
doc.add_paragraph('Nivel 3: Desarrollo (correccion de bugs, nuevas funcionalidades).', style='List Bullet')

doc.add_page_break()

# ============================================================
# 9. APENDICES
# ============================================================
doc.add_heading('9. Apendices', level=1)

doc.add_heading('9.1 Comandos Utiles', level=2)
tabla = doc.add_table(rows=9, cols=2)
tabla.style = 'Light Grid Accent 1'
cmds = [
    ('Comando', 'Descripcion'),
    ('python manage.py runserver', 'Iniciar servidor de desarrollo'),
    ('python manage.py migrate', 'Ejecutar migraciones de Django'),
    ('python manage.py createsuperuser', 'Crear superusuario administrador'),
    ('python crear_superusuario.py', 'Crear superusuario del proyecto'),
    ('npm run build', 'Compilar frontend Vue.js'),
    ('npm run dev', 'Compilar frontend en modo watch'),
    ('python scripts/asegurar_tipos_movimiento.py', 'Verificar tipos de movimiento en BD'),
    ('python scripts/validate_schema.py', 'Validar esquema contra modelos Django'),
]
for i, row_data in enumerate(cmds):
    for j, val in enumerate(row_data):
        tabla.rows[i].cells[j].text = val
        if i == 0:
            for p in tabla.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('9.2 Triggers de Base de Datos', level=2)
doc.add_paragraph(
    'El proyecto incluye triggers SQL en el directorio scripts/ que controlan '
    'la gestion de stock. Ejecutarlos despues de importar el esquema de la base de datos:'
)

doc.add_paragraph('trigger_proteccion_stock.sql: Trigger consolidado con proteccion contra stock negativo. Reemplaza trg_actualizar_stock_oferta.', style='List Bullet')
doc.add_paragraph('trigger_stock_vendida.sql: Trigger que descuenta stock cuando un movimiento cambia a estado vendida.', style='List Bullet')
doc.add_paragraph('trigger_modificar_stock.sql: Version anterior del trigger (reemplazada por proteccion_stock).', style='List Bullet')

doc.add_heading('9.3 Estructura de Directorios', level=2)
doc.add_paragraph(
    'agrosft/\n'
    '  config/          - Configuracion del proyecto Django\n'
    '  core/            - Clases base, middleware, helpers compartidos\n'
    '  apps/\n'
    '    usuarios/      - Autenticacion, registro, panel admin\n'
    '    inventario/    - Catalogo de productos y marketplace\n'
    '    ventas/        - Carrito, solicitudes, ventas, calificaciones\n'
    '    clientes/      - Historial de clientes\n'
    '    facturacion/   - Facturacion y generacion de PDF\n'
    '  frontend/        - Componentes Vue.js\n'
    '  templates/       - Plantillas HTML globales\n'
    '  static/          - Archivos estaticos compilados\n'
    '  media/           - Archivos subidos por usuarios\n'
    '  scripts/         - Scripts SQL y utilidades'
)

doc.add_page_break()

# ============================================================
# 10. GLOSARIO
# ============================================================
doc.add_heading('10. Glosario de Terminos', level=1)
glosario = [
    ('Backend', 'Parte del servidor de una aplicacion que procesa logica de negocio y acceso a datos.'),
    ('CSRF', 'Cross-Site Request Forgery. Proteccion contra ataques de falsificacion de solicitudes.'),
    ('Django', 'Framework web de alto nivel para Python.'),
    ('Entorno virtual', 'Aislamiento de dependencias de Python por proyecto (venv).'),
    ('Frontend', 'Parte visual de la aplicacion con la que interactua el usuario.'),
    ('Managed', 'Atributo de modelos Django que indica si la BD gestiona las migraciones.'),
    ('Marketplace', 'Plataforma donde se exhiben y venden productos de multiples vendedores.'),
    ('Migracion', 'Archivo que define los cambios en el esquema de la base de datos.'),
    ('Middleware', 'Componente que procesa las peticiones HTTP antes de llegar a la vista.'),
    ('Movimiento', 'Transaccion de compra o venta registrada en el sistema.'),
    ('OAuth', 'Protocolo de autorizacion para autenticacion con servicios de terceros.'),
    ('Pipeline', 'Secuencia de pasos en la autenticacion social para procesar datos del usuario.'),
    ('Trigger', 'Procedimiento almacenado en MySQL que se ejecuta automaticamente ante eventos.'),
    ('Venv', 'Modulo de Python para crear entornos virtuales.'),
    ('Vue.js', 'Framework JavaScript progresivo para construir interfaces de usuario.'),
]
tabla = doc.add_table(rows=len(glosario)+1, cols=2)
tabla.style = 'Light Grid Accent 1'
tabla.rows[0].cells[0].text = 'Termino'
tabla.rows[0].cells[1].text = 'Definicion'
for p in tabla.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in tabla.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for i, (term, defi) in enumerate(glosario):
    tabla.rows[i+1].cells[0].text = term
    tabla.rows[i+1].cells[1].text = defi

doc.add_page_break()

# ============================================================
# 11. BIBLIOGRAFIA
# ============================================================
doc.add_heading('11. Bibliografia', level=1)
refs = [
    'Django Software Foundation. (2024). Django 5.0 Documentation. https://docs.djangoproject.com/',
    'Vue.js. (2024). Vue.js 3 Documentation. https://vuejs.org/guide/introduction.html',
    'MySQL. (2024). MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/',
    'Vite. (2024). Vite Documentation. https://vitejs.dev/guide/',
    'Bootstrap. (2024). Bootstrap 5.1 Documentation. https://getbootstrap.com/docs/5.1/',
    'SENA. (2024). Repositorio oficial del SENA. Referencia de manual tecnico.',
    'Brevo. (2024). Brevo API Documentation. https://developers.brevo.com/',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_page_break()

# ============================================================
# 12. INDICE ANALITICO
# ============================================================
doc.add_heading('12. Indice Analitico', level=1)
indices = [
    ('A', 'Asistencia tecnica, Autenticacion, Auditoria'),
    ('B', 'Base de datos, Bootstrap, Brevo'),
    ('C', 'Carrito de compras, Configuracion, CSFR'),
    ('D', 'Desinstalacion, Django, .env'),
    ('E', 'Ejecucion, Entorno virtual, Esquema'),
    ('F', 'Facturacion, Frontend'),
    ('G', 'Garantias, Git, Google OAuth'),
    ('H', 'Hardware (requerimientos)'),
    ('I', 'Instalacion, Inventario'),
    ('L', 'Licenciamiento'),
    ('M', 'Marketplace, Migraciones, Movimiento, MySQL'),
    ('N', 'Node.js, NPM'),
    ('P', 'Panel de administracion, Permisos, Productos'),
    ('R', 'Requerimientos, Repositorio'),
    ('S', 'SENA, Seguridad, Sesion, Software (requerimientos)'),
    ('T', 'Tabla de movimientos, Triggers'),
    ('U', 'Usuarios'),
    ('V', 'Validacion de esquema, Ventas, Vue.js'),
]
for letra, terminos in indices:
    p = doc.add_paragraph()
    r = p.add_run(f'{letra}: ')
    r.bold = True
    p.add_run(terminos)

# ============================================================
# GUARDAR
# ============================================================
output_path = r'C:\Users\samup\OneDrive\Documentos\SENA\PORTAFOLIO\agrosft\docs\Manual_de_Instalacion_AgroSFT.docx'
doc.save(output_path)
print(f'Manual generado exitosamente: {output_path}')
